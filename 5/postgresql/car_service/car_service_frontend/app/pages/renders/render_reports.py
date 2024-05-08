import datetime
import os
import tempfile

import streamlit as st
import altair as alt
import pandas as pd
from docx import Document
from docx.shared import Inches


def foreign_domestic_cars_service_cost(conn):
    st.markdown("## Общая стоимость обслуживания отечественных и импортных автомобилей")
    today = datetime.datetime.now()
    month_ago = today - datetime.timedelta(days=30)
    date = st.date_input(
        "Выберите период обслуживания машин",
        (month_ago, today + datetime.timedelta(days=1)),
        max_value=today + datetime.timedelta(days=1),
        format="YYYY-MM-DD",
    )
    if date and len(date) == 2:
        total_service_cost_query = f"""
        SELECT
            cars.is_foreign,
            SUM(
                CASE WHEN cars.is_foreign THEN
                    services.cost_foreign
                ELSE
                    services.cost_our
                END) AS service_cost
        FROM
            works
            LEFT JOIN cars ON works.car_id = cars.id
            LEFT JOIN services ON works.service_id = services.id
            WHERE date_work >= TO_DATE('{date[0]}','YYYY-MM-DD') 
            AND date_work <= TO_DATE('{date[1]}', 'YYYY-MM-DD')
        GROUP BY
            cars.is_foreign
        ORDER BY
            service_cost DESC;
        """
        if st.checkbox("Использовать хранимую функцию", key="service-cost-sql-funciton-checkbox"):
            total_service_cost_query = f"""SELECT * from service_cost_by_foreignness('{date[0]}','{date[1]}')"""
        with st.expander("Исполняемый SQL"):
            st.code(total_service_cost_query, "sql")
        total_service_cost = conn.query(total_service_cost_query)
        c = (
            alt.Chart(total_service_cost)
            .mark_arc(outerRadius=80)
            .encode(theta=alt.Theta("service_cost:Q").stack(True), color=alt.Color("is_foreign:N").legend(None))
        )
        text = c.mark_text(radius=125, size=20).encode(
            text=alt.condition(alt.datum.is_foreign, alt.value("Импорт"), alt.value("Отечеств."))
        )
        text_val = c.mark_text(radius=100, size=20).encode(text="service_cost:Q")
        c = c + text + text_val
        c = c.properties(title=f'Стоимость обслуживания по "импортности" автомобилей [{date[0]} - {date[1]}] (в RUB)')
        # c = c.configure_title(fontSize=20, anchor='middle')
        st.altair_chart(c, use_container_width=True)
        st.dataframe(total_service_cost, hide_index=True)
        return c, total_service_cost
    else:
        st.markdown("Ожидается ввод периода...")
        return None


def best_monthly_workers(conn):
    st.markdown(f"## 5 Мастеров, выполнивших наибольшое кол-во работ за период")
    date = st.date_input("Выберите месяц работы мастеров", max_value=datetime.datetime.now(), format="YYYY-MM-DD")
    if date:
        d = date.replace(day=1).strftime("%Y-%m-%d"), (date.replace(day=1) + datetime.timedelta(days=30)).strftime(
            "%Y-%m-%d"
        )
        query = f"""
        SELECT 
            masters.id as master_id, 
            masters.name as master_name, 
            COUNT(1) as works_count 
        FROM works 
            LEFT JOIN masters ON works.master_id = masters.id 
            WHERE works.date_work >= TO_DATE('{d[0]}','YYYY-MM-DD') 
            AND works.date_work <= TO_DATE('{d[1]}', 'YYYY-MM-DD') 
        GROUP BY masters.id 
        ORDER BY 
            COUNT(1) DESC 
        LIMIT 5;"""
        if st.checkbox("Использовать хранимую функцию"):
            query = f"""SELECT * from top_five_masters('{d[0]}','{d[1]}')"""
        best_workers = conn.query(query)
        best_workers = best_workers.set_index("master_id").sort_values(
            by=[
                "works_count",
            ],
            ascending=False,
        )
        with st.expander("Исполняемый SQL"):
            st.code(query, "sql")
        c = (
            alt.Chart(best_workers)
            .encode(
                x=alt.X("works_count", title="Кол-во выполненных работ за период"),
                y=alt.Y("master_name", title=["Имя мастера"]),
                text="works_count",
            )
            .properties(title=f"Топ 5 мастеров по кол-ву работ [{d[0]} - {d[1]}]")
        )
        c = c.mark_bar() + c.mark_text(align="left", dx=2)
        st.altair_chart(c, use_container_width=True)
        st.dataframe(best_workers)
        return c, best_workers
    else:
        st.markdown("Ожидается ввод месяца...")
        return None


def add_dataframe(df: pd.DataFrame, document) -> None:
    t = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])


def generate_report_docx(service_cost_data, best_workers_data):
    """Produces docx report and return the object"""
    document = Document()

    # 1. Extract fetched data
    sc_chart, sc_df = service_cost_data
    bw_chart, bw_df = best_workers_data

    doc = b""
    # 2. Save artefacts to tempdir
    with tempfile.TemporaryDirectory() as tmpdir:
        sc_chart_path = os.path.join(tmpdir, "sc_chart.png")
        sc_chart.save(sc_chart_path)
        bw_chart_path = os.path.join(tmpdir, "bw_chart.png")
        bw_chart.save(bw_chart_path)

        document.add_heading(f'Отчёт о работе от {datetime.datetime.now().strftime("%Y-%m-%d")}', 0)
        p = document.add_paragraph(
            "Здесь приведена информация о работе сервиса. Если график пустой - то и информации нет"
        )

        if sc_df.shape[0] != 0:
            document.add_heading("Общая стоимость обслуживания отечественных и импортных автомобилей", level=1)
            document.add_picture(sc_chart_path, width=Inches(7))
            add_dataframe(sc_df, document)

        if bw_df.shape[0] != 0:
            document.add_heading("Топ 5 мастеров по кол-ву работ", level=1)
            document.add_picture(bw_chart_path, width=Inches(7))
            add_dataframe(bw_df, document)

        document_path = os.path.join(tmpdir, "report.docx")
        document.save(document_path)
        with open(document_path, "rb") as report_fb:
            doc = report_fb.read()
    return doc


def create_report(service_cost_data, best_workers_data):
    """Generates report and gives user the pdf file"""
    doc = generate_report_docx(service_cost_data, best_workers_data)
    return doc


def render_reports() -> None:
    conn = st.experimental_connection("car_service_db")
    service_cost_data = foreign_domestic_cars_service_cost(conn)
    st.markdown("---")
    best_workers_data = best_monthly_workers(conn)
    if None not in (service_cost_data, best_workers_data):
        if st.button(label="Создать отчёт", key="download-report-btn"):
            report = create_report(service_cost_data, best_workers_data)
            st.markdown("**Отчет готов!**")
            st.download_button("Скачать отчёт", report, file_name=f"{datetime.datetime.now()}_car_service_report.docx")
